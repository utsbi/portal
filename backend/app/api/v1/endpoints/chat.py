from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Request
import asyncio
from datetime import datetime

from app.schemas.chat import ChatRequest, ChatResponse, ChatMessage, SourceDocument
from app.agents.explore import run_explore_agent
from app.api.deps import get_current_user_id
from app.services.pdf_parser import PDFParser


router = APIRouter()


@router.post("/", response_model=ChatResponse)
async def chat(request: ChatRequest, raw_request: Request, user_id: str = Depends(get_current_user_id)):
    """Chat with the Explore AI Agent."""
    try:
        history = [
            {
                "role": msg.role,
                "content": msg.content
            }
            for msg in request.history
        ]

        # Run agent as a cancellable task
        agent_task = asyncio.create_task(
            run_explore_agent(
                query=request.query,
                client_id=user_id,
                history=history,
                attachments=[
                    {"filename": att.filename, "content": att.content, "file_type": att.file_type}
                    for att in request.attachments
                ],
                model_preference=request.model_preference or "fast"
            )
        )

        # Monitor for client disconnection to cancel and save tokens
        async def monitor_disconnect():
            while True:
                if await raw_request.is_disconnected():
                    return
                await asyncio.sleep(0.5)

        disconnect_task = asyncio.create_task(monitor_disconnect())

        # Wait for either agent completion or client disconnect
        done, pending = await asyncio.wait(
            {agent_task, disconnect_task},
            return_when=asyncio.FIRST_COMPLETED
        )

        # Clean up whichever task is still pending
        for task in pending:
            task.cancel()
            try:
                await task
            except (asyncio.CancelledError, Exception):
                pass

        # If client disconnected, agent was cancelled - return empty
        if disconnect_task in done:
            return ChatResponse(answer="", sources=[], timestamp=datetime.now())

        result = agent_task.result()

        sources = []
        if request.include_sources:
            for source in result.get("sources", []):
                sources.append(SourceDocument(
                    content=source.get("content", ""),
                    filename=source.get("filename", "Unknown"),
                    page_number=source.get("page_number"),
                    relevance_score=source.get("relevance_score")
                ))

        return ChatResponse(
            answer=result.get("response", ""),
            sources=sources,
            timestamp=datetime.now()
        )

    except asyncio.CancelledError:
        return ChatResponse(answer="", sources=[], timestamp=datetime.now())
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error processing chat request: {str(e)}"
        )


@router.get("/health")
async def chat_health():
    """Health check endpoint for chat service."""
    return {
        "status": "healthy",
        "service": "chat",
        "timestamp": datetime.now().isoformat()
    }


@router.post("/extract-text")
async def extract_file_text(
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user_id)
):
    """Extract text from file without adding to database.

    Supports PDF, DOCX, DOC, and TXT files.
    Returns the extracted text content for session-only use.
    """
    try:
        file_bytes = await file.read()
        filename = file.filename or "attachment"
        file_lower = filename.lower()

        if file_lower.endswith('.pdf'):
            # Use PDF parser for PDF files
            pdf_parser = PDFParser()
            pages = pdf_parser.extract_text_with_metadata(file_bytes, filename)
            content = "\n\n".join([p["content"] for p in pages])
            file_type = "pdf"

        elif file_lower.endswith(('.doc', '.docx')):
            # Use python-docx for DOCX files
            import docx
            import io
            doc = docx.Document(io.BytesIO(file_bytes))
            content = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            file_type = "docx"

        else:
            # Plain text files
            content = file_bytes.decode('utf-8', errors='ignore')
            file_type = "txt"

        return {
            "filename": filename,
            "content": content,
            "file_type": file_type
        }

    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to extract text from file: {str(e)}"
        )
